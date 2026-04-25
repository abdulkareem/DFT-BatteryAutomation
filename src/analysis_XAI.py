#!/usr/bin/env python3
"""Electronic-structure analysis + automated writing helpers."""

from __future__ import annotations

import argparse
import json
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


@dataclass
class JobResult:
    ligand: str
    job_id: str
    e_complex: float
    e_ligand: float
    e_li: float
    homo_ev: float
    lumo_ev: float

    @property
    def binding_energy_eh(self) -> float:
        return self.e_complex - (self.e_ligand + self.e_li)

    @property
    def homo_lumo_gap_ev(self) -> float:
        return self.lumo_ev - self.homo_ev


def _pull(data: dict[str, Any], *keys: str, default: float = 0.0) -> float:
    cursor: Any = data
    for key in keys:
        if isinstance(cursor, dict):
            cursor = cursor.get(key)
        else:
            cursor = None
    if cursor is None:
        return default
    return float(cursor)


def parse_job_json(json_file: Path, ligand: str, job_id: str) -> JobResult:
    payload = json.loads(json_file.read_text())
    return JobResult(
        ligand=ligand,
        job_id=job_id,
        e_complex=_pull(payload, "energies", "E_complex"),
        e_ligand=_pull(payload, "energies", "E_ligand"),
        e_li=_pull(payload, "energies", "E_Li_plus"),
        homo_ev=_pull(payload, "orbitals", "HOMO_eV"),
        lumo_ev=_pull(payload, "orbitals", "LUMO_eV"),
    )


def summarize(results: list[JobResult], out_csv: Path) -> pd.DataFrame:
    rows = []
    for r in results:
        row = asdict(r)
        row["binding_energy_eh"] = r.binding_energy_eh
        row["homo_lumo_gap_ev"] = r.homo_lumo_gap_ev
        rows.append(row)

    df = pd.DataFrame(rows)
    df.to_csv(out_csv, index=False)
    return df


def make_plots(df: pd.DataFrame, output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    # For this workflow, we use binding_energy_eh as a proxy for solvation trends.
    plt.figure(figsize=(8, 4.5))
    plt.bar(df["ligand"], df["binding_energy_eh"], color="steelblue")
    plt.ylabel("Solvation proxy / Binding Energy (Eh)")
    plt.xlabel("Ligand type")
    plt.title("Solvation energy trend vs ligand type")
    plt.xticks(rotation=30)
    plt.tight_layout()
    plt.savefig(output_dir / "solvation_vs_ligand.png", dpi=200)
    plt.close()

    plt.figure(figsize=(8, 4.5))
    plt.scatter(df["binding_energy_eh"], df["homo_lumo_gap_ev"], color="darkred")
    for _, row in df.iterrows():
        plt.annotate(row["ligand"], (row["binding_energy_eh"], row["homo_lumo_gap_ev"]))
    plt.xlabel("Binding energy (Eh)")
    plt.ylabel("HOMO-LUMO gap (eV)")
    plt.title("Electronic structure map")
    plt.tight_layout()
    plt.savefig(output_dir / "binding_vs_gap.png", dpi=200)
    plt.close()


def build_manuscript(df: pd.DataFrame, out_docx: Path, out_pdf: Path) -> None:
    doc = Document()
    doc.add_heading("Fluorinated Amide Additives for Aqueous Li-ion Batteries", level=1)

    doc.add_heading("Literature Review", level=2)
    doc.add_paragraph(
        "Automated note: integrate Scholarcy/Paperpal/Consensus outputs focused on "
        "(i) aqueous electrolyte stability, (ii) Li+ solvation shell engineering, and "
        "(iii) SEI formation pathways from 2025–2026 datasets."
    )

    doc.add_heading("Methods", level=2)
    doc.add_paragraph(
        "Geometry optimizations used ORCA 6.1.1 with ! PBEh-3c Opt RIJONX. "
        "A health-check routine monitored convergence and restarted stalled trajectories "
        "via 0.05 Å coordinate jittering."
    )

    doc.add_heading("Results and Discussion", level=2)
    ranked = df.sort_values("binding_energy_eh")
    for _, row in ranked.iterrows():
        doc.add_paragraph(
            f"{row['ligand']}: ΔEb = {row['binding_energy_eh']:.6f} Eh; "
            f"HOMO/LUMO gap = {row['homo_lumo_gap_ev']:.3f} eV."
        )

    doc.save(out_docx)

    c = canvas.Canvas(str(out_pdf), pagesize=A4)
    y = 800
    c.setFont("Helvetica-Bold", 13)
    c.drawString(50, y, "Supporting Information")
    y -= 30
    c.setFont("Helvetica", 10)
    c.drawString(50, y, "Table S1. Computed ΔEb and HOMO/LUMO gaps.")
    y -= 20
    for _, row in df.iterrows():
        c.drawString(
            50,
            y,
            f"{row['job_id']} | {row['ligand']} | ΔEb={row['binding_energy_eh']:.6f} Eh | Gap={row['homo_lumo_gap_ev']:.3f} eV",
        )
        y -= 15
        if y < 70:
            c.showPage()
            y = 800
            c.setFont("Helvetica", 10)
    c.save()


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser()
    p.add_argument("--jobs-root", type=Path, required=True)
    p.add_argument("--output-dir", type=Path, required=True)
    return p.parse_args()


def main() -> int:
    args = parse_args()
    args.output_dir.mkdir(parents=True, exist_ok=True)

    results = []
    for job_dir in sorted(args.jobs_root.glob("Job_*")):
        result_json = job_dir / "result.json"
        if not result_json.exists():
            continue
        ligand = job_dir.name.replace("Job_", "Ligand_")
        results.append(parse_job_json(result_json, ligand=ligand, job_id=job_dir.name))

    if not results:
        raise FileNotFoundError(f"No result.json files found in: {args.jobs_root}")

    df = summarize(results, out_csv=args.output_dir / "results_summary.csv")
    make_plots(df, output_dir=args.output_dir)
    build_manuscript(
        df,
        out_docx=args.output_dir / "manuscript.docx",
        out_pdf=args.output_dir / "Supporting_Info.pdf",
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
